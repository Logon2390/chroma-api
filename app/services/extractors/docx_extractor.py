import os
from pathlib import Path
from typing import Dict, Any
from datetime import datetime

import docx

from app.services.extractors.base import BaseExtractor
from app.models.schemas import FileType


class DocxExtractor(BaseExtractor):
    """Extractor for DOCX files.
    
    This class provides functionality to extract text and metadata
    from DOCX files using python-docx.
    """
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            The extracted text content as a string.
            
        Raises:
            ValueError: If the DOCX file cannot be processed.
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return "\n".join(full_text)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            A dictionary containing metadata about the DOCX file.
            
        Raises:
            ValueError: If the DOCX file cannot be processed.
        """
        try:
            metadata = {
                "filename": file_path.name,
                "file_size": os.path.getsize(file_path),
                "file_type": FileType.DOCX,
                "extracted_at": datetime.now().isoformat(),
            }
            
            # Extract document properties
            doc = docx.Document(file_path)
            core_properties = doc.core_properties
            
            if core_properties:
                docx_metadata = {}
                
                # Extract common metadata fields
                if core_properties.author:
                    docx_metadata["author"] = core_properties.author
                if core_properties.title:
                    docx_metadata["title"] = core_properties.title
                if core_properties.created:
                    docx_metadata["created"] = core_properties.created.isoformat()
                if core_properties.modified:
                    docx_metadata["modified"] = core_properties.modified.isoformat()
                if core_properties.comments:
                    docx_metadata["comments"] = core_properties.comments
                
                metadata["docx_metadata"] = docx_metadata
            
            # Count paragraphs, tables, etc.
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            
            return metadata
        except Exception as e:
            raise ValueError(f"Failed to extract metadata from DOCX: {str(e)}")
    
    @classmethod
    def supports_filetype(cls, file_type: FileType) -> bool:
        """Check if this extractor supports the given file type.
        
        Args:
            file_type: The type of file to check support for.
            
        Returns:
            True if this extractor supports DOCX files, False otherwise.
        """
        return file_type == FileType.DOCX 